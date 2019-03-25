import csv
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

#ROWS = 334
#COLS = 6

ROWS = 122
COLS = 6

#CSV = 'FR_final-1.csv'
CSV = 'FR22.csv'

with open(CSV, 'r') as csvfile:
    data = list(csv.reader(csvfile, delimiter=','))
print("Data length :", len(data))

doc = Document('report.docx')
table = doc.tables[0]
#table.allow_autofit = True

for row in range(1,ROWS):
    for col in range(COLS):
        cell = table.rows[row].cells[col]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run()
        if col == 5:
            try:
                image_path = '/home/ashwin/classifier-work/newimg-1/' + str(data[row][1]) + '_' + str(data[row][2]) + '.jpeg'
                #print(str(data[row][1]) + '_' + str(data[row][2]))
                r.add_picture(image_path, width=Cm(7.0), height=Cm(7.0))
                r.add_text('\n')
            except:
                print(str(data[row][1]) + '_' + str(data[row][2]))
        else:
            r.add_text(data[row][col])
#doc.save('test.docx')
doc.save('FRII.docx')
"""
p = tables[0].rows[1].cells[5].add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run()
r.add_picture('marsmap.jpg',width=Cm(7.0), height=Cm(7.0))
p = tables[0].rows[2].cells[5].add_paragraph()
r = p.add_run()
r.add_picture('marsmap.jpg',width=Cm(7.0), height=Cm(7.0))
doc.save('test.docx')
"""
